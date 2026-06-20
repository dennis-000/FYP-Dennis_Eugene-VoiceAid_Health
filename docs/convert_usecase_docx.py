import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor

def create_docx(md_path, docx_path):
    doc = Document()
    
    # Configure 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_mermaid = False
    
    for line in lines:
        stripped = line.rstrip()
        
        if stripped.startswith('```mermaid'):
            in_mermaid = True
            p = doc.add_paragraph()
            run = p.add_run("[Mermaid Use Case Diagram Code - Copy the block below to mermaid.live to download as PNG/SVG image]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            continue
        elif stripped.startswith('```') and in_mermaid:
            in_mermaid = False
            continue
            
        if in_mermaid:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line.rstrip('\n'))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.0)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue
            
        # Parse titles
        if stripped.startswith('# '):
            h = doc.add_heading(level=1)
            run = h.add_run(stripped[2:])
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00) # African Kente red
        elif stripped.startswith('## '):
            h = doc.add_heading(level=2)
            run = h.add_run(stripped[3:])
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif stripped.startswith('### '):
            h = doc.add_heading(level=3)
            run = h.add_run(stripped[4:])
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif stripped.startswith('* ') or stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            content = stripped[2:]
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
        elif stripped == '':
            continue
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
                    
    doc.save(docx_path)
    print(f"Successfully compiled {docx_path}")

if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.abspath(__file__))
    md = os.path.join(base_dir, 'Use_Case_Diagram.md')
    docx = os.path.join(base_dir, 'Use_Case_Diagram.docx')
    create_docx(md, docx)
