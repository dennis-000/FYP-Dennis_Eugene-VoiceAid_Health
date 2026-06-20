import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_docx(md_path, docx_path):
    doc = Document()
    
    # Configure 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default style to Calibri
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    def add_code_block(doc, code_text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.05
        
        # Add background shading to the paragraph using XML manipulation
        pPr = p._p.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA"/>')
        pPr.append(shd)
        
        # Add a subtle left border to emulate code container styling
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="8" w:color="CCCCCC"/></w:pBdr>')
        pPr.append(pBdr)
        
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.0)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        return p

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code = False
    code_lines = []
    
    for line in lines:
        stripped = line.rstrip()
        
        # Check code block toggle
        if stripped.startswith('```'):
            if in_code:
                # Ending code block
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code = False
            else:
                # Starting code block
                in_code = True
            continue
            
        if in_code:
            code_lines.append(line.rstrip('\r\n'))
            continue
            
        # Parse titles and headers
        if stripped.startswith('# '):
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(stripped[2:])
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00) # Deep Red (Kente)
        elif stripped.startswith('## '):
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(stripped[3:])
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif stripped.startswith('### '):
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(stripped[4:])
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        # Bullet list
        elif stripped.startswith('* ') or stripped.startswith('- '):
            content = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            # Bold parser (**text**)
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        elif stripped == '---':
            doc.add_page_break()
        elif stripped == '':
            continue
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.abspath(__file__))
    md = os.path.join(base_dir, 'Appendix_A_Source_Code_Listings.md')
    docx = os.path.join(base_dir, 'Appendix_A_Source_Code_Listings.docx')
    create_docx(md, docx)
